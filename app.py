import io
from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.oxml.ns import qn

app = Flask(__name__)

def process_paragraph_runs(p, replacements=None):
    """
    Identifies each contiguous block of highlighted runs independently,
    preventing separate highlighted fields from merging across lines.
    """
    highlights = []
    runs = p.runs
    i = 0

    while i < len(runs):
        run = runs[i]
        rPr = run._r.get_or_add_rPr()
        highlight_elem = rPr.find(qn('w:highlight'))

        if highlight_elem is not None:
            # Collect contiguous highlighted runs for a SINGLE field
            block_runs = [run]
            color = highlight_elem.get(qn('w:val')) or "yellow"
            i += 1

            while i < len(runs):
                next_run = runs[i]
                next_rPr = next_run._r.get_or_add_rPr()
                next_highlight = next_rPr.find(qn('w:highlight'))
                
                # Check if contiguous and matching highlight
                if next_highlight is not None:
                    block_runs.append(next_run)
                    i += 1
                else:
                    break

            # Combine text ONLY for this specific block
            block_text = "".join([r.text for r in block_runs]).strip()

            if block_text:
                highlights.append({"text": block_text, "color": color})

            # If replacing values for output document
            if replacements is not None:
                if block_text in replacements and replacements[block_text] != "":
                    new_val = replacements[block_text]
                    
                    first_run = block_runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    bold = first_run.bold
                    italic = first_run.italic

                    # Replace text in first run of the block only
                    first_run.text = new_val

                    # Restore font properties
                    if font_name: first_run.font.name = font_name
                    if font_size: first_run.font.size = font_size
                    first_run.bold = bold
                    first_run.italic = italic

                    # Clear remaining runs in this block
                    for r in block_runs[1:]:
                        r.text = ""

                # Remove highlight XML from all runs in this block
                for r in block_runs:
                    r_rPr = r._r.get_or_add_rPr()
                    h_elem = r_rPr.find(qn('w:highlight'))
                    if h_elem is not None:
                        r_rPr.remove(h_elem)
        else:
            i += 1

    return highlights

def process_document(doc, replacements=None):
    highlights = []

    # Process paragraphs
    for p in doc.paragraphs:
        highlights.extend(process_paragraph_runs(p, replacements))

    # Process table cells
    for table in doc.tables:
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    highlights.extend(process_paragraph_runs(p, replacements))

    return highlights

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/parse-doc', methods=['POST'])
def parse_doc():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    doc = Document(file)
    highlights = process_document(doc)
    return jsonify({"highlights": highlights})

@app.route('/generate-doc', methods=['POST'])
def generate_doc():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    replacements = request.form.to_dict()

    doc = Document(file)
    process_document(doc, replacements)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="Updated_Toll_Receipt.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == '__main__':
    app.run(debug=True, port=5000)